# #integrations/rag_app/main.py
# import os
# import sys
# import json
# import logging
# import tempfile

# import numpy as np
# import boto3
# import psycopg2
# import PyPDF2
# import pandas as pd
# import ssl

# from PIL import Image
# from flask import Flask, request, jsonify
# from flask_cors import CORS
# from moviepy import AudioFileClip
# from dotenv import load_dotenv

# # ensure repo root on path so we can import tools.py
# sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
# from tools import (
#     reload_rag_index,
#     get_db_connection,
#     _rag_search,
#     embed_text_or_image,
# )

# load_dotenv()
# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)

# # Bedrock & Whisper configs
# BEDROCK_REGION        = os.getenv("AWS_REGION", "us-east-1")
# AWS_ACCESS_KEY        = os.getenv("AWS_ACCESS_KEY")
# AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
# MODEL_ID_EMBED        = os.getenv("MODEL_ID_EMBED")
# MODEL_ID_CHAT         = os.getenv("MODEL_ID_CHAT")

# # Lazy-load Whisper
# _whisper_model = None
# def get_whisper_model():
#     global _whisper_model
#     if _whisper_model is None:
#         import whisper
#         _whisper_model = whisper.load_model("base")
#         logger.info("Whisper model loaded")
#     return _whisper_model

# # Bedrock client (for image/text embed on upload)
# bedrock = boto3.client(
#     "bedrock-runtime",
#     region_name=BEDROCK_REGION,
#     aws_access_key_id=AWS_ACCESS_KEY,
#     aws_secret_access_key=AWS_SECRET_ACCESS_KEY
# )
# logger.info("Bedrock client initialized.")

# app = Flask(__name__)
# CORS(app)

# # preload the RAG index once at startup
# reload_rag_index()
# logger.info("Initial RAG index loaded")

# def split_text(text, max_words=2000):
#     words = text.split()
#     for i in range(0, len(words), max_words):
#         yield " ".join(words[i : i + max_words])

# def transcribe_audio(path):
#     return get_whisper_model().transcribe(path)["text"]

# def extract_text(file):
#     filename = file.filename.lower()
#     suffix = os.path.splitext(filename)[1]
#     fd, tmp = tempfile.mkstemp(suffix=suffix)
#     os.close(fd)
#     file.save(tmp)
#     logger.info(f"Saved temp file at {tmp}")

#     try:
#         if suffix in [".mp3", ".wav", ".m4a"]:
#             return {"text": transcribe_audio(tmp), "image_base64": None, "filename": filename}

#         if suffix in [".mp4", ".mov", ".mkv"]:
#             aud = tmp + "_audio.wav"
#             clip = AudioFileClip(tmp)
#             clip.write_audiofile(aud)
#             clip.close()
#             return {"text": transcribe_audio(aud), "image_base64": None, "filename": filename}

#         if suffix in [".jpg", ".jpeg", ".png"]:
#             with open(tmp, "rb") as f:
#                 img_bytes = bytearray(f.read())
#             tex = boto3.client(
#                 "textract",
#                 region_name=BEDROCK_REGION,
#                 aws_access_key_id=AWS_ACCESS_KEY,
#                 aws_secret_access_key=AWS_SECRET_ACCESS_KEY
#             )
#             blocks = tex.detect_document_text(Document={"Bytes": img_bytes})["Blocks"]
#             lines = [b["Text"] for b in blocks if b["BlockType"] == "LINE"]
#             return {"text": "\n".join(lines), "image_base64": None, "filename": filename}

#         if suffix == ".pdf":
#             reader = PyPDF2.PdfReader(tmp)
#             text = "\n".join(p.extract_text() or "" for p in reader.pages)
#             return {"text": text, "image_base64": None, "filename": filename}

#         if suffix == ".docx":
#             import docx  # pick up python-docx, not any stray docx.py
#             doc = docx.Document(tmp)
#             texts = [p.text for p in doc.paragraphs if p.text.strip()]
#             for tbl in doc.tables:
#                 for row in tbl.rows:
#                     texts.append(" | ".join(c.text for c in row.cells if c.text.strip()))
#             return {"text": "\n".join(texts), "image_base64": None, "filename": filename}

#         if suffix in [".txt", ".py", ".csv", ".xlsx"]:
#             if suffix == ".csv":
#                 df = pd.read_csv(tmp)
#                 return {"text": df.to_string(index=False), "image_base64": None, "filename": filename}
#             if suffix == ".xlsx":
#                 df = pd.read_excel(tmp, engine="openpyxl")
#                 return {"text": df.to_string(index=False), "image_base64": None, "filename": filename}
#             return {"text": open(tmp, encoding="utf-8").read(), "image_base64": None, "filename": filename}

#         raise ValueError("Unsupported file type")
#     finally:
#         try: os.remove(tmp)
#         except: pass

# @app.route("/upload", methods=["POST"])
# def upload_file():
#     if "file" not in request.files:
#         return jsonify({"error": "No file uploaded"}), 400

#     file = request.files["file"]
#     try:
#         ext = extract_text(file)
#     except Exception as e:
#         logger.error(f"Extraction failed: {e}")
#         return jsonify({"error": "Extraction failed"}), 500

#     text = ext["text"]
#     img_b64 = ext.get("image_base64")
#     fn = ext["filename"]

#     conn, cur = get_db_connection()
#     cur.execute("""
#       CREATE TABLE IF NOT EXISTS embeddings (
#         id SERIAL PRIMARY KEY,
#         file_name TEXT,
#         chunk_index INTEGER,
#         embedding FLOAT8[],
#         embedding_size INTEGER,
#         chunk_text TEXT
#       )
#     """); conn.commit()

#     count = 0
#     for i, chunk in enumerate(split_text(text), start=1):
#         emb = embed_text_or_image(chunk, content_type="text", model_id=MODEL_ID_EMBED)
#         cur.execute(
#             "INSERT INTO embeddings (file_name,chunk_index,embedding,embedding_size,chunk_text) VALUES (%s,%s,%s,%s,%s)",
#             (fn, i, emb, len(emb), chunk)
#         )
#         count += 1

#     if img_b64:
#         try:
#             emb = embed_text_or_image(img_b64, content_type="image", model_id=MODEL_ID_CHAT)
#             cur.execute(
#                 "INSERT INTO embeddings (file_name,chunk_index,embedding,embedding_size,chunk_text) VALUES (%s,%s,%s,%s,%s)",
#                 (fn, 9999, emb, len(emb), "[IMAGE]")
#             )
#             count += 1
#         except Exception:
#             logger.warning("Image embedding failed")

#     conn.commit()
#     cur.close()
#     conn.close()

#     reload_rag_index()
#     logger.info("RAG index reloaded after upload")

#     return jsonify({"message": f"Embedded {count} chunks"})

# @app.route("/chat", methods=["POST"])
# def chat():
#     data = request.get_json()
#     q = data.get("message") or data.get("question")
#     if not q:
#         return jsonify({"error": "Missing question"}), 400

#     # single-shot RAG search + answer
#     answer = _rag_search(q)
#     return jsonify({"answer": answer})

# @app.route("/voice-chat", methods=["POST"])
# def start_voice_console():
#     return jsonify({
#       "message": "To start a voice-to-voice console session, run:\n"
#                  "    python agent.py console"
#     })

# if __name__ == "__main__":
#     app.run(host="0.0.0.0", port=5000, debug=True)














# integrations/rag_app/main.py

import os
import sys
import logging
import tempfile

from flask import Flask, request, jsonify
from flask_cors import CORS
from moviepy import AudioFileClip
import PyPDF2
import pandas as pd
import boto3
from dotenv import load_dotenv

# ensure we can import tools.py from repo root
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))
from tools import reload_rag_index, embed_text_or_image, get_db_connection, _rag_search

load_dotenv()
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# AWS/Textract config (only needed for OCR on images)
AWS_REGION = os.getenv("AWS_REGION", "us-east-1")
AWS_ACCESS_KEY = os.getenv("AWS_ACCESS_KEY")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
textract = boto3.client(
    "textract",
    region_name=AWS_REGION,
    aws_access_key_id=AWS_ACCESS_KEY,
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY
)

app = Flask(__name__)
CORS(app)

# permanent storage for uploaded files
data_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "data"))
os.makedirs(data_dir, exist_ok=True)

# ─── Helpers ─────────────────────────────────────────────────────────────

def split_text(text: str, max_words: int = 2000):
    words = text.split()
    for i in range(0, len(words), max_words):
        yield " ".join(words[i : i + max_words])

_whisper_model = None
def get_whisper_model():
    global _whisper_model
    if _whisper_model is None:
        import whisper
        _whisper_model = whisper.load_model("base")
        logger.info("Whisper model loaded")
    return _whisper_model

def transcribe_audio(path: str) -> str:
    return get_whisper_model().transcribe(path)["text"]

def extract_text_from(path: str) -> dict:
    """
    Given a filesystem path, detect file type and extract its text.
    Returns a dict: {"text": ..., "image_base64": None, "filename": basename}
    """
    filename = os.path.basename(path).lower()
    suffix = os.path.splitext(filename)[1]
    # Audio
    if suffix in [".mp3", ".wav", ".m4a"]:
        return {"text": transcribe_audio(path), "image_base64": None, "filename": filename}

    # Video
    if suffix in [".mp4", ".mov", ".mkv"]:
        wav = path + "_audio.wav"
        clip = AudioFileClip(path)
        clip.write_audiofile(wav)
        clip.close()
        text = transcribe_audio(wav)
        os.remove(wav)
        return {"text": text, "image_base64": None, "filename": filename}

    # Image → Textract OCR
    if suffix in [".jpg", ".jpeg", ".png"]:
        with open(path, "rb") as f:
            img_bytes = f.read()
        resp = textract.detect_document_text(Document={"Bytes": img_bytes})
        lines = [b["Text"] for b in resp["Blocks"] if b["BlockType"] == "LINE"]
        return {"text": "\n".join(lines), "image_base64": None, "filename": filename}

    # PDF
    if suffix == ".pdf":
        reader = PyPDF2.PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return {"text": "\n".join(pages), "image_base64": None, "filename": filename}

    # Word
    if suffix == ".docx":
        import docx  # ensure python-docx
        doc = docx.Document(path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                texts.append(" | ".join(c.text for c in row.cells if c.text.strip()))
        return {"text": "\n".join(texts), "image_base64": None, "filename": filename}

    # CSV / XLSX / TXT / PY
    if suffix in [".csv", ".xlsx", ".txt", ".py"]:
        if suffix == ".csv":
            df = pd.read_csv(path)
            return {"text": df.to_string(index=False), "image_base64": None, "filename": filename}
        if suffix == ".xlsx":
            df = pd.read_excel(path, engine="openpyxl")
            return {"text": df.to_string(index=False), "image_base64": None, "filename": filename}
        # plain text / code
        with open(path, "r", encoding="utf-8") as f:
            return {"text": f.read(), "image_base64": None, "filename": filename}

    raise ValueError(f"Unsupported file type: {suffix}")

def bootstrap_embeddings():
    """
    On startup, process any files in data_dir that aren't yet in Postgres.
    """
    conn, cur = get_db_connection()
    cur.execute("SELECT DISTINCT file_name FROM embeddings")
    seen = {r[0] for r in cur.fetchall()}
    for fname in os.listdir(data_dir):
        if fname in seen:
            continue
        path = os.path.join(data_dir, fname)
        try:
            data = extract_text_from(path)
        except Exception as e:
            logger.warning(f"Skipping {fname}: extraction failed: {e}")
            continue
        text = data["text"]
        for idx, chunk in enumerate(split_text(text), start=1):
            emb = embed_text_or_image(chunk, content_type="text")
            cur.execute(
                "INSERT INTO embeddings (file_name, chunk_index, embedding, embedding_size, chunk_text) VALUES (%s,%s,%s,%s,%s)",
                (fname, idx, emb, len(emb), chunk)
            )
        conn.commit()
        logger.info(f"Bootstrapped embeddings for {fname}")
    cur.close()
    conn.close()

# ─── Startup ───────────────────────────────────────────────────────────────

bootstrap_embeddings()
reload_rag_index()
logger.info("Initial data_dir embeddings loaded and RAG index built")

# ─── Routes ────────────────────────────────────────────────────────────────

@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    perm_path = os.path.join(data_dir, file.filename)
    file.save(perm_path)
    logger.info(f"Saved permanent copy: {perm_path}")

    try:
        data = extract_text_from(perm_path)
    except Exception as e:
        logger.error(f"Extraction failed: {e}")
        return jsonify({"error": "Extraction failed"}), 500

    text, fn = data["text"], data["filename"]
    conn, cur = get_db_connection()
    cur.execute("""
      CREATE TABLE IF NOT EXISTS embeddings (
        id SERIAL PRIMARY KEY,
        file_name TEXT,
        chunk_index INTEGER,
        embedding FLOAT8[],
        embedding_size INTEGER,
        chunk_text TEXT
      );
    """)
    conn.commit()

    count = 0
    for idx, chunk in enumerate(split_text(text), start=1):
        emb = embed_text_or_image(chunk, content_type="text")
        cur.execute(
            "INSERT INTO embeddings (file_name, chunk_index, embedding, embedding_size, chunk_text) VALUES (%s,%s,%s,%s,%s)",
            (fn, idx, emb, len(emb), chunk)
        )
        count += 1

    conn.commit()
    cur.close()
    conn.close()

    reload_rag_index()
    logger.info("RAG index reloaded after upload")

    return jsonify({"message": f"Embedded {count} chunks from {fn}"})

@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json() or {}
    q = data.get("message") or data.get("question")
    if not q:
        return jsonify({"error": "Missing question"}), 400
    answer = _rag_search(q)
    return jsonify({"answer": answer})

@app.route("/voice-chat", methods=["POST"])
def start_voice_console():
    return jsonify({
        "message": "To start a voice-to-voice console session, run:\n"
                   "    python agent.py console"
    })

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)




